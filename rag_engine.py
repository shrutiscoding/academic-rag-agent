# rag_engine.py
# ============================================================
# FINAL FIXED VERSION
# ✔ Correct file upload to Supabase Storage
# ✔ Proper MIME types (PDF/DOCX/PPTX/CSV)
# ✔ Download works perfectly
# ✔ Existing RAG features preserved
# ============================================================

import os
import re
import uuid
import tempfile
import warnings
import requests
import pandas as pd
import pdfplumber
import mimetypes
import pytesseract

import pytesseract
from PIL import Image
import io


from docx import Document as DocxDocument

from difflib import get_close_matches
from pathlib import Path
from dotenv import load_dotenv
from supabase import create_client

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer

warnings.filterwarnings("ignore")
import os

if os.name == "nt":
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# ============================================================
# ENV
# ============================================================

load_dotenv()

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")

supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

CHUNK_SIZE = 1200
CHUNK_OVERLAP = 150


# ============================================================
# EMBEDDINGS
# ============================================================

class Embeddings:
    def __init__(self):
        self.model = SentenceTransformer(
            "all-MiniLM-L6-v2",
            device="cpu"
        )

    def embed(self, text):
        if not text:
            text = "empty"
        return self.model.encode(text).tolist()


# ============================================================
# ENGINE
# ============================================================

class RAGEngine:

    def __init__(self):

        self.embedder = Embeddings()

        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP
        )

    # ========================================================
    # LLM
    # ========================================================

    def ask_llm(self, prompt):

        try:
            r = requests.post(
                OPENROUTER_URL,
                headers={
                    "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "openai/gpt-4o-mini",
                    "messages": [
                        {"role": "user", "content": prompt}
                    ]
                },
                timeout=120
            )

            data = r.json()

            if "choices" in data:
                return data["choices"][0]["message"]["content"].strip()

            return "No response."

        except Exception as e:
            return str(e)
        
    def clean_text(self, text):
        if not text:
            return ""

        text = text.replace("\x00", "")   # fix DB error
        text = re.sub(r'\s+', ' ', text).strip()

        return text

    # ========================================================
    # LANGUAGE DETECT
    # ========================================================

    def detect_language(self, text):

        try:
            prompt = f"""
Detect language of text.

Return only one word:
English
Hindi
Marathi
Mixed

Text:
{text}
"""
            x = self.ask_llm(prompt).strip()

            if x in ["English", "Hindi", "Marathi", "Mixed"]:
                return x

            return "English"

        except:
            return "English"

    # ========================================================
    # STORAGE FIXED
    # ========================================================

    def upload_file_to_storage(self, file_bytes, filename):

        try:
            unique = f"{uuid.uuid4()}_{filename}"

            mime_type, _ = mimetypes.guess_type(filename)

            if not mime_type:
                mime_type = "application/octet-stream"

            supabase.storage.from_("documents").upload(
                path=unique,
                file=file_bytes,
                file_options={
                    "content-type": mime_type,
                    "upsert": "true"
                }
            )

            return unique

        except Exception as e:
            print("Upload Error:", e)
            return None

    def get_download_url(self, filename):

        try:
            files = supabase.storage.from_("documents").list()

            for f in files:
                if f["name"].endswith("_" + filename):
                    return supabase.storage.from_("documents").get_public_url(
                        f["name"]
                    )

            return None

        except:
            return None

    # ========================================================
    # FILE LOADERS
    # ========================================================
    def load_pdf(self, file_bytes):

        docs = []

        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(file_bytes)
            path = tmp.name

        try:
            with pdfplumber.open(path) as pdf:

                for i, page in enumerate(pdf.pages):

                    text = page.extract_text()

                    # ✅ OCR WITHOUT POPPLER
                    if not text or len(text.strip()) < 20:
                        try:
                            img = page.to_image(resolution=300).original
                            text = pytesseract.image_to_string(img)
                        except Exception as e:
                            print("OCR error:", e)
                            text = ""

                    text = self.clean_text(text)

                    if text:
                        docs.append(
                            Document(
                                page_content=text,
                                metadata={"page": i + 1}
                            )
                        )

        finally:
            os.remove(path)

        return docs
            

    def load_docx(self, file_bytes):

        docs = []

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(file_bytes)
            path = tmp.name

        try:
            doc = DocxDocument(path)

            text = "\n".join([p.text for p in doc.paragraphs])
            text = self.clean_text(text)

            if text:
                docs.append(
                    Document(
                        page_content=text,
                        metadata={"page": 1}
                    )
                )

            return docs

        finally:
            os.remove(path)

    def load_text(self, file_bytes):

        text = file_bytes.decode(
            "utf-8",
            errors="ignore"
        )

        return [
            Document(
                page_content=text,
                metadata={"page": 1}
            )
        ]

    def load_csv(self, file_bytes):

        tmp = tempfile.NamedTemporaryFile(
            delete=False,
            suffix=".csv"
        )

        try:
            tmp.write(file_bytes)
            tmp.close()

            df = pd.read_csv(tmp.name)

            return [
                Document(
                    page_content=df.to_string(),
                    metadata={"page": 1}
                )
            ]

        finally:
            os.remove(tmp.name)

    def load_file(self, file_bytes, filename):

        ext = Path(filename).suffix.lower()

        if ext == ".pdf":
            return self.load_pdf(file_bytes)
        elif ext == ".docx":
            return self.load_docx(file_bytes)

        elif ext == ".csv":
            return self.load_csv(file_bytes)
        
        elif ext in [".png", ".jpg", ".jpeg"]:
            return self.load_image(file_bytes)

        else:
            return self.load_text(file_bytes)

    # ========================================================
    # SUMMARY / TERMS
    # ========================================================

    def generate_summary(self, text):
        return self.ask_llm(
            f"Summarize:\n{text[:12000]}"
        )

    def extract_key_terms(self, text):

        r = self.ask_llm(
            f"""
Extract 15 keywords comma separated only.

{text[:8000]}
"""
        )

        arr = [
            x.strip()
            for x in r.replace("\n", ",").split(",")
            if x.strip()
        ]

        return list(dict.fromkeys(arr))[:15]

    # ========================================================
    # UPLOAD PROCESS
    # ========================================================

    def process_upload(self, file_bytes, filename):

    # Check if file already exists
        existing = supabase.table("documents") \
            .select("id") \
            .eq("filename", filename) \
            .execute()

        if existing.data:
            return 0   # already uploaded

        # upload to storage
        self.upload_file_to_storage(file_bytes, filename)

        docs = self.load_file(file_bytes, filename)

        chunks = self.splitter.split_documents(docs)

        full = self.clean_text(" ".join(d.page_content for d in docs))

        summary = self.generate_summary(full)

        key_terms = ", ".join(self.extract_key_terms(full))

        ins = supabase.table("documents").insert({
            "filename": filename,
            "file_type": Path(filename).suffix.replace(".", "").upper(),
            "chunks": len(chunks),
            "summary": summary,
            "key_terms": key_terms
        }).execute()

        doc_id = ins.data[0]["id"]

        rows = []

        for i, chunk in enumerate(chunks):

            clean_chunk = self.clean_text(chunk.page_content)

            rows.append({
                "document_id": doc_id,
                "content": clean_chunk,
                "embedding": self.embedder.embed(clean_chunk),
                "page": chunk.metadata.get("page", 1),
                "chunk_no": i + 1
            })

        supabase.table("document_chunks").insert(rows).execute()

        return len(chunks)

    # ========================================================
    # DOCUMENTS
    # ========================================================

    def get_all_documents(self):

        try:
            r = supabase.table("documents").select("*").order(
                "uploaded_date",
                desc=True
            ).execute()

            return r.data or []

        except:
            return []

    # ========================================================
    # SEARCH / CHAT SAME
    # ========================================================

    def normalize_question(self, question):

        words = re.findall(r'\w+', question.lower())

        db_words = []

        try:
            rows = supabase.table("documents").select(
                "key_terms"
            ).execute().data

            for row in rows:
                if row["key_terms"]:
                    db_words.extend(
                        row["key_terms"].lower().split(",")
                    )

        except:
            pass

        final = []

        for w in words:

            if len(w) <= 2:
                final.append(w)
                continue

            match = get_close_matches(
                w,
                db_words,
                n=1,
                cutoff=0.75
            )

            if match:
                final.append(match[0].strip())
            else:
                final.append(w)

        return " ".join(final)

    # ========================================================
    # KEYWORD SEARCH
    # ========================================================

    def keyword_search(self, question):

        words = re.findall(r'\w+', question.lower())

        words = [w for w in words if len(w) > 2][:6]

        results = []

        for word in words:

            try:

                r = supabase.table("document_chunks") \
                    .select("content,page,document_id") \
                    .ilike("content", f"%{word}%") \
                    .limit(5) \
                    .execute()

                results.extend(r.data)

            except:
                pass

        return results
    


    def load_image(self, file_bytes):

        try:
            image = Image.open(io.BytesIO(file_bytes))

            text = pytesseract.image_to_string(image)

            text = self.clean_text(text)

            return [
                Document(
                    page_content=text,
                    metadata={"page": 1}
                )
            ]

        except Exception as e:
            print("Image OCR error:", e)
            return []

    # ========================================================
    # CHAT
    # ========================================================

    def answer_question(self, question):

        try:

            fixed_question = self.normalize_question(question)

            qvec = self.embedder.embed(fixed_question)

            vr = supabase.rpc(
                "match_documents",
                {
                    "query_embedding": qvec,
                    "match_count": 20
                }
            ).execute()

            vector_rows = vr.data or []

            keyword_rows = self.keyword_search(fixed_question)

            rows = vector_rows + keyword_rows

            if not rows:
                return "No documents found.", [], False

            q_words = set(
                w.lower()
                for w in re.findall(r'\w+', fixed_question)
                if len(w) > 2
            )

            ranked = []

            for row in rows:

                txt = row["content"].lower()

                overlap = sum(
                    1 for w in q_words if w in txt
                )

                sim = float(row.get("similarity", 0.30))

                score = sim + (overlap * 0.30)

                if score >= 0.35:
                    ranked.append((score, row))

            if not ranked:
                return "Information not found in uploaded documents.", [], False

            ranked.sort(reverse=True, key=lambda x: x[0])

            top_rows = [x[1] for x in ranked[:3]]

            context = "\n\n".join(
                r["content"] for r in top_rows
            )

            lang = self.detect_language(question)

            prompt = f"""
Use ONLY context below.

Reply in same language user asked.
Language = {lang}

If unavailable say:
Information not found in uploaded documents.

Context:
{context}

Question:
{question}
"""

            answer = self.ask_llm(prompt)

            if "information not found" in answer.lower():
                return answer, [], False

            # =================================================
            # REAL SOURCE FILTER
            # =================================================

            sources = []
            used = set()

            answer_words = set(
                w.lower()
                for w in re.findall(r'\w+', answer)
                if len(w) > 3
            )

            for row in top_rows:

                txt = row["content"].lower()

                q_match = sum(
                    1 for w in q_words if w in txt
                )

                a_match = sum(
                    1 for w in answer_words if w in txt
                )

                if q_match == 0:
                    continue

                if a_match == 0 and q_match < 2:
                    continue

                filename = row.get("filename")

                if not filename:

                    doc_id = row["document_id"]

                    d = supabase.table("documents") \
                        .select("filename") \
                        .eq("id", doc_id) \
                        .limit(1) \
                        .execute()

                    if d.data:
                        filename = d.data[0]["filename"]

                page = row.get("page", 1)

                key = f"{filename}_{page}"

                if key not in used:

                    used.add(key)

                    sources.append({
                        "filename": filename,
                        "page": page,
                        "url": self.get_download_url(filename)
                    })

            return answer, sources, True

        except Exception as e:
            return f"Error: {str(e)}", [], False